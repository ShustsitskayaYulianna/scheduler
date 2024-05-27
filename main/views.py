from django.shortcuts import render
import requests
from django.http import HttpResponseRedirect
from django.template import RequestContext

from .forms import DateForm, TimeForm
from requests_html import HTMLSession
import time
import pandas as pd
from bs4 import BeautifulSoup
import re
import datetime
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt
from django.http import HttpResponse

def index(request):
    # if this is a POST request we need to process the form data
    if request.method == "POST":
        # create a form instance and populate it with data from the request:
        form = DateForm(request.POST)
        # check whether it's valid:
        if form.is_valid():
            # process the data in form.cleaned_data as required
            # ...
            # redirect to a new URL:
            return HttpResponseRedirect("")

    # if a GET (or any other method) we'll create a blank form
    else:
        form = DateForm()

    return render(request, 'main/index.html', {"form": form})


def date(request):
    context = dict()
    if request.method == 'POST':
        form = TimeForm(request.POST)
        only_all_dates_list = list()
        only_all_days_list = list()
        session = HTMLSession()
        holidays2 = list()
        after_check_holidays_list = list()
        all_data_in_list = list()
        only_month_in_list = list()
        all_data_and_holiday = list()
        len_list = list()
        worship_list = list()
        time_list = list()
        mon_worship = list()
        links = list()


        #print(request.POST)
        date_request = request.POST["date"]
        try:
            valid_date = time.strptime(date_request, '%Y-%m-%d')
            dates_and_days_from_system = pd.DataFrame({
                "Date": pd.date_range(start=date_request, periods=pd.Period(date_request).days_in_month).date,
                "Day": pd.date_range(start=date_request, periods=pd.Period(date_request).days_in_month).day_name(),
            })

            # get data from collumn date in DataFrame
            all_date_from_collum = dates_and_days_from_system['Date'].tolist()
            for only_all_dates in all_date_from_collum:
                coll = str(only_all_dates)
                coll.split(' ', 1)[0]
                only_all_dates_list.append(coll)

            all_day_from_collum = dates_and_days_from_system['Day'].tolist()
            for only_all_month in all_day_from_collum:
                month = str(only_all_month)
                if month == "Monday":
                    month_RU = "Понедельник"
                elif month == "Tuesday":
                    month_RU = "Вторник"
                elif month == "Wednesday":
                    month_RU = "Среда"
                elif month == "Thursday":
                    month_RU = "Четверг"
                elif month == "Friday":
                    month_RU = "Пятница"
                elif month == "Saturday":
                    month_RU = "Суббота"
                elif month == "Sunday":
                    month_RU = "Воскресение"
                only_all_days_list.append(month_RU)

        except ValueError:
            print('Неверный формат даты. Необходимо ввести в виде год-месяц-день (Например: 2022-01-10)')

        for date in only_all_dates_list:
            saint_list = list()
            saint_list_title = list()
            holids_list_title = list()

            bogosluz_ukaz = 'https://azbyka.ru/bogosluzhebnye-ukazaniya?date=' + date

            azbuka_days = 'https://azbyka.ru/days/' + date
            links.append(azbuka_days)
            k = session.get(azbuka_days)

            API_URL = 'https://azbyka.ru/days/api/day/' + date + '.json'

            API_response = requests.get(API_URL).json()
            API_holidays = API_response["holidays"]
            for API_holids in API_holidays:
                API_holids = API_holids["title"]
                holids_list_title.append(API_holids)

            API_saints = API_response["saints"]
            for API_saint in API_saints:
                API_saint = API_saint["title_genitive"]
                saint_list_title.append(API_saint)

            i = 0
            if len(holids_list_title) == 0:
                if len(saint_list_title) > 3:
                    saint_list.append(saint_list_title[0])
                    saint_list.append(saint_list_title[1])
                    saint_list.append(saint_list_title[2])
                elif len(saint_list_title) < 3:
                    for i in range(len(saint_list_title)):
                        saint_list.append(saint_list_title[i])
                elif len(saint_list_title) == 0:
                    for i in range(len(holids_list_title)):
                        saint_list.append(holids_list_title[i])
            else:
                for i in range(len(holids_list_title)):
                    saint_list.append(holids_list_title[i])
                if len(saint_list_title) > 1:
                    saint_list.append(saint_list_title[0])
                    saint_list.append(saint_list_title[1])
                if len(saint_list_title) > 2:
                    saint_list.append(saint_list_title[2])
                else:
                    print("")

            #print("saint_list", saint_list)

            page = requests.get(azbuka_days)  # - BeautifulSoup
            soup = BeautifulSoup(page.text, "html.parser")

            week_of_day = list()
            svetlaya_sedmica = k.html.xpath('//*[@class="shadow"]/b')

            pominovenie_1 = k.html.xpath('//*[@class="shadow"]/a')
            for p in pominovenie_1:
                if "Поминовение усопших" in p.text:
                    week_of_day.append(p)

            all_holid_with_week = soup.findAll('a', href="https://azbyka.ru/nedelya")
            for all in all_holid_with_week:
                if str(all.text) == "Неделя":
                    if "о мытаре и фарисее" in str(all.nextSibling.text):
                        week_of_day.append(all)
                        week_of_day.append(all.nextSibling)
                    elif "о блудном сыне" in str(all.nextSibling.text):
                        week_of_day.append(all)
                        week_of_day.append(all.nextSibling)

                    else:
                        print("")
                else:
                    print("")

            all_holid_with_family = soup.findAll('a', href="https://azbyka.ru/roditelskie-subboty")
            for all in all_holid_with_family:
                week_of_day.append(all)

            kanon = soup.findAll('a',
                                 href="https://azbyka.ru/molitvoslov/velikij-kanon-svt-andreya-kritskogo-s-perevodom-na-russkij-yazyk.html")
            for all in kanon:
                week_of_day.append(all)

            all_holid_about_mytare_i_farisee = soup.findAll('a',
                                                            href="https://azbyka.ru/days/nedelja-o-mytare-i-farisee")
            for all in all_holid_about_mytare_i_farisee:
                week_of_day.append(all)

            all_holid_with_subday = soup.findAll('a',
                                                 href="/days/prazdnik-nedelja-syropustnaja-vospominanie-adamova-izgnanija-proshchenoe-voskresene")
            for all in all_holid_with_subday:
                week_of_day.append(all)
                week_of_day.append(all.nextSibling)

            all_holid_with_vel_post_pominovenie = soup.findAll('a', href="/velikiy-post")
            for all in all_holid_with_vel_post_pominovenie:
                if "Поминовение усопших":
                    week_of_day.append(all.nextSibling)
                else:
                    print("")

            all_holid_with_torzestvo = soup.findAll('a', href="https://azbyka.ru/days/prazdnik-torzhestvo-pravoslavija")
            for all in all_holid_with_torzestvo:
                week_of_day.append(all)

            all_holid_bogorodicy = soup.findAll('a',
                                                href="/days/prazdnik-pohvala-presvjatoj-bogorodicy-subbota-akafista")
            for all in all_holid_bogorodicy:
                week_of_day.append(all)

            all_holid_hrista = soup.findAll('a',
                                            href="/days/prazdnik-velikaja-pjatnica-vospominanie-svjatyh-spasitelnyh-strastej-gospoda-nashego-iisusa-hrista")
            for all in all_holid_hrista:
                week_of_day.append(all)

            all_holid_with_krestopoklonnaja = soup.findAll('a',
                                                           href="/days/prazdnik-nedelja-3-ja-velikogo-posta-krestopoklonnaja")
            for all in all_holid_with_krestopoklonnaja:
                week_of_day.append(all)

            all_holid_with_sb = soup.findAll('a', href="/days/prazdnik-lazareva-subbota")
            for all in all_holid_with_sb:
                week_of_day.append(all)

            all_holid_svyatych_otec = soup.findAll('a', href="/days/nedelja-svjatyh-otec")
            for all in all_holid_svyatych_otec:
                week_of_day.append(all)

            all_holid_with_strastnaja_sedmica = soup.findAll('a', href="/days/p-strastnaja-sedmica")
            for all in all_holid_with_strastnaja_sedmica:
                week_of_day.append(all)
                if (str(all.nextSibling) == '. '):
                    week_of_day.append(all.nextSibling.nextSibling)
                else:
                    week_of_day.append(all.nextSibling)

            for all in svetlaya_sedmica:
                week_of_day.append(all)

            all_holid_with_radonica = soup.findAll('a', href="https://azbyka.ru/days/prazdnik-radonica")
            for all in all_holid_with_radonica:
                if "Радоница" in str(all.text):
                    week_of_day.append(all)
                    week_of_day.append(all.nextSibling)

            all_holid_with_antipasxa = soup.findAll('a', href="https://azbyka.ru/antipasxa")
            for all in all_holid_with_antipasxa:
                week_of_day.append(all)

            all_holid_with_jen = soup.findAll('a', href="https://azbyka.ru/days/sv-marija-magdalina-mironosica")
            for all in all_holid_with_jen:
                week_of_day.append(all)

            all_holid_with_war = soup.findAll('a', href="/days/prazdnik-pominovenie-usopshih-voinov")
            for all in all_holid_with_war:
                week_of_day.append(all)

            all_holid_with_war = soup.findAll('a', href="https://azbyka.ru/days/prazdnik-sobor-vseh-svjatyh")
            for all in all_holid_with_war:
                week_of_day.append(all)

            all_holid_with_war = soup.findAll('a',
                                              href="https://azbyka.ru/days/sv-sobor-vseh-svjatyh-v-zemle-rossijskoj-prosijavshih")
            for all in all_holid_with_war:
                week_of_day.append(all)

            all_holid_with_dmitr_sub = soup.findAll('a', href="/days/prazdnik-dimitrievskaja-roditelskaja-subbota")
            for all in all_holid_with_dmitr_sub:
                week_of_day.append(all)

            all_holid_with_dmitr_sub = soup.findAll('a', href="/days/prazdnik-svjatyh-praotec")
            for all in all_holid_with_dmitr_sub:
                week_of_day.append(all)

            all_holidays = list()

            hol1 = k.html.xpath('//*[@class="text day__text"]/p/a')

            a = list()
            holid = ""
            for h in hol1:
                for s in saint_list:
                    s = re.sub(r" \(.*?\)", "", s)
                    if s in h.text:
                        # if h == h:
                        a.append(h)
                        # holid +=  re.sub(r" \(.*?\)", "", h.text)
                        # h = h.text
                        # if h.text == h.text:

            # print(all_holidays)
            for one_student_choice in a:
                if one_student_choice not in all_holidays:
                    all_holidays.append(one_student_choice)

            for all in week_of_day:
                holid += str(all.text).strip() + ". "
                # hol = hol.strip() + ". "
                holid = re.sub(', Крестопоклонная ', 'Неделя Крестопоклонная. ', holid)
                holid = re.sub('Крестопоклонная ', 'Неделя Крестопоклонная. ', holid)
                holid = re.sub(', Крестопоклонная. ', 'Неделя Крестопоклонная. ', holid)
                holid = re.sub(r'\( ', '', holid)
                holid = re.sub(r'\(.', '', holid)
                holid = re.sub(r' . .', '', holid)
                # holid = re.sub('Светлая', '', holid)
                # holid = re.sub('сплошная.', '', holid)
                # holid = re.sub('–', '', holid)
                # all_holidays.append(holid)

            for all in all_holidays:
                # holid +=  re.sub(r" \(.*?\)", "", all) + ". "
                # holid += all.text
                holid += all.text.strip() + ". "
                if holid == holid:
                    holid = holid
                    holid = re.sub(r" \(\d+\)", " ", holid)
                    holid = re.sub(r" \(ок. \d+\)", " ", holid)  # -удаляет (ок. 400
                    holid = re.sub(r" \(ок. .*?\)", " ", holid)
                    holid = re.sub('седмица', '', holid)
                    holid = re.sub('Страстна́я .', '', holid)
                # holid = re.sub('..', '', holid)

            # print(all_holidays)
            # print(holid)

            data_without_year = date.replace('-', ' ').split(' ', 1)[1]
            only_month = data_without_year.split(' ', 1)[0]
            only_data = data_without_year.split(' ', 1)[1]
            year = date.replace('-', ' ').split(' ', 1)[0]

            if only_month == "10":
                month_number = 10
            elif only_month == "11":
                month_number = 11
            elif only_month == "12":
                month_number = 12
            else:
                for number in only_month:
                    month_number = int(number)

            month = ""
            only_month = datetime.date(1900, month_number, 1).strftime('%B')
            if only_month == "January":
                month = "января"
            elif only_month == "February":
                month = "февраля"
            elif only_month == "March":
                month = "марта"
            elif only_month == "April":
                month = "апреля"
            elif only_month == "May":
                month = "мая"
            elif only_month == "June":
                month = "июня"
            elif only_month == "July":
                month = "июля"
            elif only_month == "August":
                month == "августа"
            elif only_month == "September":
                month == "сентября"
            elif only_month == "October":
                month = "октября"
            elif only_month == "November":
                month = "ноября"
            elif only_month == "December":
                month = "декабря"

            all_data_in_list.append(only_data)
            only_month_in_list.append(month)
            try:
                all_data_and_holiday.append(' ' + holid)
            except IndexError:
                print("")

            len_hol = len(all_data_and_holiday)
            len_list = list(range(1, len_hol + 1))

        for i in links:
            print(i)

        context = {'all_data': zip(all_data_and_holiday, only_all_days_list, all_data_in_list, len_list, links),
                   'month': month,
                   'year': year,
                   'worship': worship_list,
                   'time': time_list
                   }
    else:
        form = DateForm()
    return render(request, 'main/name.html', context)


def view(request):
    len_list = list();

    days = request.GET.getlist('days')
    holidays = request.GET.getlist('holidays')
    month = request.GET.get('month')
    year = request.GET.get('year')
    worships = request.GET.getlist('outputWorship')

    len_hol = len(holidays)
    len_list = list(range(1, len_hol + 1))

    for i in range(len(holidays)):
        worshipsDiv = request.GET.getlist('worshipsDiv-'+str(i))

    return render(request, 'main/view.html', {'days_and_holidays': zip(days, holidays, worships),
                                              'month': month,
                                              'year': year})

def createDoc(request):
    days = request.GET.getlist('days')
    month = request.GET.get('month')
    year = request.GET.get('year')
    holidays = request.GET.getlist('holidays')
    worships = request.GET.getlist('outputWorship')
    worshipslist = list()
    worshipslistMap = list()

    for worship in worships:
        worship = worship.replace('\r\n', '\n')
        worship = worship[:-1]
        worshipslist.append(worship)


    header = "РАСПИСАНИЕ БОГОСЛУЖЕНИЙ В ПРИХОДЕ ХРАМА ПРЕОБРАЖЕНИЯ ГОСПОДНЯ г. МИНСКА на " + month + " " + year + " г."

    doc_main = docx.Document()
    main_header = doc_main.add_paragraph()
    main_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_in_bold = main_header.add_run(header).bold = True
    for i in range(len(holidays)):
        data_doc = days
        data_hol = holidays
        par1 = doc_main.add_paragraph()
        par1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        par1_bold_italic = par1.add_run(data_doc[i])
        par1_bold_italic = par1.add_run(data_hol[i])
        par1_bold_italic.italic = True
        par1.paragraph_format.space_before = Pt(3)
        par1.paragraph_format.space_after = Pt(1)
        p_fmt = par1.paragraph_format
        p_fmt.left_indent = Cm(1)

        par2 = doc_main.add_paragraph(worshipslist[i])
        par2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_fmt = par2.paragraph_format
        p_fmt.left_indent = Cm(3)

    for paragraph in doc_main.paragraphs:
        style = doc_main.styles['Normal']
        font = style.font
        font.date = 'Times New Roman'
        font.size = Pt(15)

    sections = doc_main.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(0.1)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    file_format = 'Расписание месяц %s.docx' % month

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename = "download.docx"'
    doc_main.save(response)
    return response


    return response

    #return HttpResponse(doc_main.save(file_format))